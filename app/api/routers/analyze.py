# app/api/routers/analyze.py

import os
from pathlib import Path
import shutil
from typing import List

from fastapi import APIRouter, HTTPException
from fastapi.responses import JSONResponse, FileResponse
from app.core.config import UPLOAD_DIRECTORY, ANALYZE_DIRECTORY
from docx import Document # <--- IMPORT INI

router = APIRouter(
    prefix="/analyze",
    tags=["Analysis"]
)

@router.post("/{file_name}", status_code=200)
async def analyze_repository(file_name: str):
    repo_file_path = UPLOAD_DIRECTORY / file_name

    if not repo_file_path.exists() or not repo_file_path.is_file():
        raise HTTPException(status_code=404, detail=f"File repositori '{file_name}' tidak ditemukan.")

    # --- SIMULASI PROSES ANALISIS YANG MEMAKAN WAKTU ---
    import time
    time.sleep(3) # Simulasi proses yang lama (misal 3 detik)

    # --- SIMULASI HASIL ANALISIS ---
    analysis_id = f"analysis_{Path(file_name).stem}_{int(time.time())}"
    
    RESULTS_DIRECTORY = ANALYZE_DIRECTORY
    RESULTS_DIRECTORY.mkdir(parents=True, exist_ok=True) 

    graph_json_filename = f"{analysis_id}_dependency_graph.json"
    doc_filename = f"{analysis_id}_documentation.docx" 

    # CREATE GRAPH JSON FILE
    graph_path = RESULTS_DIRECTORY / graph_json_filename
    
    # --- Membuat file DOCX dummy yang VALID dengan python-docx ---
    dummy_doc_path = RESULTS_DIRECTORY / doc_filename
    document = Document()
    document.add_heading(f'Repository Documentation for {file_name}', 0)
    document.add_paragraph(f'This is a generated documentation for the repository file: {file_name}.')
    
    document.add_heading('Analysis Summary', level=1)
    document.add_paragraph(f'Total Classes: 18')
    document.add_paragraph(f'Total Functions: 42')
    document.add_paragraph(f'Total Methods: 25')

    document.add_heading('Dependency Graph Overview', level=1)
    document.add_paragraph('A detailed dependency graph is available for download as JSON.')
    
    document.add_paragraph('Further analysis details would be presented here in a real application.')
    document.save(dummy_doc_path) # Simpan sebagai file .docx

    return JSONResponse(status_code=200, content={
        "message": f"Analisis repositori '{file_name}' berhasil.",
        "analysis_id": analysis_id,
        "total_classes": 18, 
        "total_functions": 42,
        "total_methods": 25,
        "graph_json_download_url": f"/api/analyze/download-result/{graph_json_filename}",
        "doc_download_url": f"/api/analyze/download-result/{doc_filename}"
    })

# --- Endpoint download-result tidak berubah, sudah menangani tipe file .docx ---
@router.get("/download-result/{result_filename}")
async def download_analysis_result(result_filename: str):
    RESULTS_DIRECTORY = UPLOAD_DIRECTORY / "analysis_results"
    file_path = RESULTS_DIRECTORY / result_filename

    if not file_path.exists() or not file_path.is_file():
        raise HTTPException(status_code=404, detail=f"File hasil '{result_filename}' tidak ditemukan.")
    
    media_type = "application/octet-stream" 
    if result_filename.endswith(".json"):
        media_type = "application/json"
    elif result_filename.endswith(".png"):
        media_type = "image/png"
    elif result_filename.endswith(".svg"):
        media_type = "image/svg+xml"
    elif result_filename.endswith(".md"):
        media_type = "text/markdown"
    elif result_filename.endswith(".pdf"):
        media_type = "application/pdf"
    elif result_filename.endswith(".docx"):
        media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    return FileResponse(
        path=file_path,
        media_type=media_type,
        filename=result_filename,
    )