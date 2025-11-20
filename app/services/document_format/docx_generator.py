import os
from datetime import datetime
from typing import List, Dict, Any
import docx
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn
from app.core.config import GRAPH_VISUALIZATION_DIRECTORY

# --- KONFIGURASI BAHASA ---
TRANSLATIONS = {
    "id": {
        "title_suffix": "Dokumentasi Teknis",
        "generated_on": "Dibuat pada",
        "toc_title": "Daftar Isi",
        "type": "Tipe",
        "file": "Lokasi File",
        "params_header": "Parameters",
        "attrs_header": "Attributes",
        "returns_header": "Returns",
        "raises_header": "Raises",
        "examples_header": "Examples",
        "col_name": "Nama",
        "col_type": "Tipe",
        "col_desc": "Deskripsi",
        "back_to_toc_tooltip": "Klik untuk kembali ke Daftar Isi",
        
        "yields_header": "Yields",
        "receives_header": "Receives",
        "warns_header": "Warns",
        "warnings_sec_header": "Warnings",
        "see_also_header": "See Also",
        "notes_header": "Notes",
        "col_warning": "Warning Type",
        "col_ref": "Reference",
        
        "graph_header": "Visualisasi Dependensi", 
        "summary_header": "Summary",             
        "description_header": "Description",     
    },
    "en": {
        "title_suffix": "Technical Documentation",
        "generated_on": "Generated on",
        "toc_title": "Table of Contents",
        "type": "Type",
        "file": "File Path",
        "params_header": "Parameters",
        "attrs_header": "Attributes",
        "returns_header": "Returns",
        "raises_header": "Raises",
        "examples_header": "Examples",
        "col_name": "Name",
        "col_type": "Type",
        "col_desc": "Description",
        "back_to_toc_tooltip": "Click to return to Table of Contents",
        
        "yields_header": "Yields",
        "receives_header": "Receives",
        "warns_header": "Warns",
        "warnings_sec_header": "Warnings",
        "see_also_header": "See Also",
        "notes_header": "Notes",
        "col_warning": "Warning Type",
        "col_ref": "Reference",
        
        "graph_header": "Dependency Visualization",
        "summary_header": "Summary",             
        "description_header": "Description",     
    }
}

class DocxDocumentationGenerator:
    def __init__(self, project_name: str, language: str = "id", use_table_format: bool = True):
        self.document = Document()
        self.project_name = project_name
        self.lang_code = language if language in TRANSLATIONS else "id"
        self.labels = TRANSLATIONS[self.lang_code]
        self.use_table_format = use_table_format # <-- Fitur Baru
        self._setup_styles()
        self.TOC_BOOKMARK = "TOC_ANCHOR"

    def _setup_styles(self):
        """Mengatur style kustom."""
        styles = self.document.styles
        try:
            code_style = styles.add_style('CodeBlock', WD_STYLE_TYPE.PARAGRAPH)
            code_style.base_style = styles['Normal']
            font = code_style.font
            font.name = 'Consolas'
            font.size = Pt(9)
            font.color.rgb = RGBColor(50, 50, 50)
            p_format = code_style.paragraph_format
            p_format.left_indent = Inches(0.2)
            p_format.space_before = Pt(4)
            p_format.space_after = Pt(4)
            
            # 2. Style Signature Block
            sig_style = styles.add_style('SignatureBlock', WD_STYLE_TYPE.PARAGRAPH)
            sig_style.base_style = styles['Normal']
            sig_style.font.name = 'Consolas'
            sig_style.font.size = Pt(12)  
            sig_style.font.bold = True    
            sig_style.font.color.rgb = RGBColor(0, 0, 0)
            sig_style.paragraph_format.left_indent = Inches(0)
            
            # Style untuk Daftar Isi (Method/Indented)
            toc_method_style = styles.add_style('TOCMethod', WD_STYLE_TYPE.PARAGRAPH)
            toc_method_style.base_style = styles['Normal']
            toc_method_style.paragraph_format.left_indent = Inches(0.3) # Indentasi untuk method
        except:
            pass 

    # --- HELPER XML UNTUK HYPERLINK (Advanced) ---
    def _add_bookmark_start(self, paragraph, bookmark_name):
        """Menambahkan awal bookmark pada paragraf."""
        tag = paragraph._p
        start = docx.oxml.shared.OxmlElement('w:bookmarkStart')
        start.set(docx.oxml.ns.qn('w:id'), '0') # ID dummy, Word akan memperbaikinya
        start.set(docx.oxml.ns.qn('w:name'), bookmark_name)
        tag.append(start)

    def _add_bookmark_end(self, paragraph):
        """Menambahkan akhir bookmark pada paragraf."""
        tag = paragraph._p
        end = docx.oxml.shared.OxmlElement('w:bookmarkEnd')
        end.set(docx.oxml.ns.qn('w:id'), '0')
        tag.append(end)

    def _add_hyperlink_text(self, paragraph, text, bookmark_name, tooltip=None):
        """
        Menambahkan teks yang bisa diklik (Hyperlink internal) ke paragraf.
        """
        # Create the hyperlink tag
        hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
        hyperlink.set(docx.oxml.ns.qn('w:anchor'), bookmark_name) # Link to bookmark
        
        if tooltip:
             hyperlink.set(docx.oxml.ns.qn('w:tooltip'), tooltip)

        # Create the run tag
        new_run = docx.oxml.shared.OxmlElement('w:r')
        
        # Create the text tag
        new_text = docx.oxml.shared.OxmlElement('w:t')
        new_text.text = text
        new_run.append(new_text)
        
        # Styling agar terlihat seperti link (Biru & Underline)
        rPr = docx.oxml.shared.OxmlElement('w:rPr')
        
        color = docx.oxml.shared.OxmlElement('w:color')
        color.set(docx.oxml.ns.qn('w:val'), '0563C1') # Standard Link Blue
        rPr.append(color)
        
        u = docx.oxml.shared.OxmlElement('w:u')
        u.set(docx.oxml.ns.qn('w:val'), 'single')
        rPr.append(u)

        new_run.append(rPr)
        hyperlink.append(new_run)
        
        paragraph._p.append(hyperlink)

    # --- FITUR UTAMA ---

    def add_title_page(self):
        """Membuat halaman judul."""
        title = self.document.add_heading(self.project_name, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        subtitle = self.document.add_paragraph(self.labels["title_suffix"])
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.runs[0].font.size = Pt(16)
        subtitle.runs[0].font.color.rgb = RGBColor(100, 100, 100)

        date_str = f"{self.labels['generated_on']}: {datetime.now().strftime('%Y-%m-%d %H:%M')}"
        date_para = self.document.add_paragraph(date_str)
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def add_table_of_contents(self, components: List[Any]):
        """
        Membuat Daftar Isi Manual yang interaktif.
        """
        # 1. Buat Judul ToC & Pasang Bookmark Utama (Untuk tombol 'Back to ToC')
        toc_header = self.document.add_heading(self.labels["toc_title"], level=1)
        self._add_bookmark_start(toc_header, self.TOC_BOOKMARK)
        self._add_bookmark_end(toc_header)
        
        # 2. Loop komponen untuk membuat daftar
        for comp in components:
            # Tentukan style: Normal untuk Class/Function, Indented untuk Method
            style = 'Normal'
            prefix = ""
            
            if comp.component_type == 'method':
                style = 'TOCMethod' # Style kustom dengan indentasi
                prefix = "• "       # Bullet point visual untuk method
            
            p = self.document.add_paragraph(style=style)
            
            # Buat Text yang nge-link ke Component ID
            display_text = f"{prefix}{comp.id}"
            # Bersihkan ID untuk nama bookmark (Word bookmark tidak boleh ada spasi/karakter aneh tertentu, 
            # tapi comp.id biasanya dot notation yg aman, kecuali panjang)
            safe_bookmark = comp.id.replace(" ", "_") 
            
            self._add_hyperlink_text(p, display_text, safe_bookmark)

        self.document.add_page_break()

    # --- HELPER BARU: RENDERING LOGIC (TABLE vs TEXT) ---
    def _render_section(self, title: str, data: List[Dict], fields: List[str], headers: List[str] = None):
        """
        Helper cerdas untuk merender seksi (Parameter, Returns, dll) 
        baik sebagai Tabel atau List Teks berdasarkan self.use_table_format.
        
        Args:
            title: Judul seksi (Heading 3)
            data: List of dictionaries (isi data)
            fields: Nama key di JSON yang akan diambil (misal: ['name', 'type', 'description'])
            headers: Label header untuk Tabel/Teks (misal: ['Nama', 'Tipe', 'Deskripsi'])
        """
        if not data:
            return

        self.document.add_heading(title, level=3)

        # --- OPSI 1: FORMAT TABEL ---
        if self.use_table_format:
            # Gunakan headers jika ada, jika tidak gunakan nama fields
            col_labels = headers if headers else [f.capitalize() for f in fields]
            
            table = self.document.add_table(rows=1, cols=len(fields))
            table.style = 'Table Grid'
            
            # Header Row
            hdr_cells = table.rows[0].cells
            for i, label in enumerate(col_labels):
                hdr_cells[i].text = str(label)
                for p in hdr_cells[i].paragraphs:
                    for r in p.runs:
                        r.font.bold = True
            
            # Data Rows
            for item in data:
                row_cells = table.add_row().cells
                for i, field_key in enumerate(fields):
                    # Ambil value, handle None/missing
                    val = item.get(field_key, "") or ""
                    # Khusus field 'name' atau 'error', kita buat bold otomatis di tabel? Opsional.
                    row_cells[i].text = str(val)

        # --- OPSI 2: FORMAT TEKS (REVISI FIX) ---
        else:
            for item in data:
                p = self.document.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.2) # Indentasi item utama
                p.paragraph_format.space_after = Pt(2)
                
                # 1. Tentukan Kunci Utama (Bold/Utama) dan Kunci Tipe (Italic/Kurung)
                primary_key = None
                primary_val = ""
                type_val = ""
                
                # Prioritas penentuan kunci utama
                if 'name' in fields: 
                    primary_key = 'name'
                elif 'error' in fields: 
                    primary_key = 'error'
                elif 'warning' in fields: 
                    primary_key = 'warning'
                elif 'type' in fields: 
                    # Kasus Returns/Yields (hanya type dan desc, tidak ada name)
                    primary_key = 'type'

                # Ambil nilai utama
                if primary_key:
                    primary_val = str(item.get(primary_key, "") or "")

                # Ambil nilai tipe (jika ada, dan jika tipe bukan kunci utama)
                if 'type' in fields and primary_key != 'type':
                    type_val = str(item.get('type', "") or "")

                # 2. RENDER BAGIAN UTAMA (Nama/Error/Warning/Return Type)
                if primary_val:
                    run = p.add_run(primary_val)
                    
                    # Styling khusus untuk 'See Also' agar monospace
                    if title == self.labels["see_also_header"]:
                        run.font.name = 'Consolas'
                        run.font.size = Pt(9)
                    else:
                        # Default: Bold untuk nama parameter/error/return type
                        run.bold = True

                # 3. RENDER TIPE (Dalam kurung, Italic)
                # Contoh: user_id (str)
                if type_val:
                    p.add_run(f" ({type_val})").italic = True

                # 4. RENDER DESKRIPSI (Baris Baru & Indented)
                desc_val = str(item.get('description', "") or "")
                
                if desc_val:
                    desc_p = self.document.add_paragraph(desc_val)
                    desc_p.paragraph_format.left_indent = Inches(0.5) # Menjorok lebih dalam dari nama
                    desc_p.paragraph_format.space_after = Pt(8)       # Jarak antar item
                else:
                    # Jika tidak ada deskripsi, beri jarak pada paragraf nama
                    p.paragraph_format.space_after = Pt(8)
    
    # --- UPDATE METHOD UTAMA ---
    def add_component_documentation(self, component: Any):
        """
        Menambahkan dokumentasi komponen.
        """
        doc_data = component.docgen_final_state.get("final_state", {}).get("documentation_json", {})
        if not doc_data:
            return

        safe_bookmark_name = component.id.replace(" ", "_")

        # 1. Header & Bookmark
        h = self.document.add_heading(level=1)
        self._add_bookmark_start(h, safe_bookmark_name)
        self._add_hyperlink_text(h, component.id, self.TOC_BOOKMARK, tooltip=self.labels["back_to_toc_tooltip"])
        self._add_bookmark_end(h)
        
        # 2. Metadata
        meta_para = self.document.add_paragraph()
        meta_para.add_run(f"{self.labels['type']}: {component.component_type.capitalize()} | ").bold = True
        meta_para.add_run(f"{self.labels['file']}: {component.relative_path}").italic = True

        # --- FITUR BARU 1: Code Signature ---
        if component.component_signature:
            # Gunakan style CodeBlock
            self.document.add_paragraph(component.component_signature, style='SignatureBlock')
        
        # --- 4. DEPENDENCY GRAPH IMAGE (REVISI: HEIGHT LIMIT) ---
        graph_url = component.dependency_graph_url
        
        if graph_url:
            full_image_path = os.path.join(str(GRAPH_VISUALIZATION_DIRECTORY), graph_url)
            
            if os.path.exists(full_image_path):
                try:
                    # 1. Masukkan gambar dengan target LEBAR 6 inci dulu
                    # Kita simpan objek gambar ke variabel 'pic' untuk dimanipulasi
                    pic = self.document.add_picture(full_image_path, width=Inches(6.0))
                    
                    # 2. Tentukan Batas Tinggi Maksimal (Misal: 8 inci agar muat satu halaman A4/Letter)
                    max_height = Inches(2.5)
                    
                    # 3. Cek apakah tinggi gambar melebihi batas?
                    if pic.height > max_height:
                        # Hitung Aspect Ratio saat ini (Lebar / Tinggi)
                        aspect_ratio = pic.width / pic.height
                        
                        # Reset Tinggi ke Batas Maksimal
                        pic.height = max_height
                        
                        # Hitung ulang Lebar agar gambar tidak gepeng (maintain aspect ratio)
                        pic.width = int(max_height * aspect_ratio)
                    
                    # 4. Tengahkan Gambar (Styling)
                    last_paragraph = self.document.paragraphs[-1]
                    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                except Exception as e:
                    print(f"[DOC GEN WARN] Gagal menambahkan gambar untuk {component.id}: {e}")
        # --- END DEPENDENCY GRAPH IMAGE ---
        
        # 3. Summary
        short_sum = doc_data.get("short_summary", "")
        ext_sum = doc_data.get("extended_summary", "")
        if short_sum:
            self.document.add_heading(self.labels["summary_header"], level=3)
            self.document.add_paragraph(short_sum).bold = True
        if ext_sum:
            self.document.add_heading(self.labels["description_header"], level=3)
            self.document.add_paragraph(ext_sum)

        # --- MENGGUNAKAN HELPER BARU UNTUK SEKSI ---
        
        # 4. Parameters
        self._render_section(
            title=self.labels["params_header"],
            data=doc_data.get("parameters", []),
            fields=['name', 'type', 'description'],
            headers=[self.labels["col_name"], self.labels["col_type"], self.labels["col_desc"]]
        )

        # 5. Attributes
        self._render_section(
            title=self.labels["attrs_header"],
            data=doc_data.get("attributes", []),
            fields=['name', 'type', 'description'],
            headers=[self.labels["col_name"], self.labels["col_type"], self.labels["col_desc"]]
        )

        # 6. Returns
        # Returns di skema baru adalah List. Jika masih dict (legacy), bungkus jadi list
        returns_data = doc_data.get("returns")
        if isinstance(returns_data, dict): returns_data = [returns_data]
        
        self._render_section(
            title=self.labels["returns_header"],
            data=returns_data or [],
            fields=['type', 'description'],
            headers=[self.labels["col_type"], self.labels["col_desc"]]
        )

        # 7. Yields
        self._render_section(
            title=self.labels["yields_header"],
            data=doc_data.get("yields", []),
            fields=['type', 'description'],
            headers=[self.labels["col_type"], self.labels["col_desc"]]
        )

        # 8. Receives
        self._render_section(
            title=self.labels["receives_header"],
            data=doc_data.get("receives", []),
            fields=['name', 'type', 'description'],
            headers=[self.labels["col_name"], self.labels["col_type"], self.labels["col_desc"]]
        )

        # 9. Raises (Field di JSON: 'error', 'description')
        # Mapping field 'error' ke kolom Nama agar konsisten dengan logic render
        self._render_section(
            title=self.labels["raises_header"],
            data=doc_data.get("raises", []),
            fields=['error', 'description'], 
            headers=[self.labels["col_name"], self.labels["col_desc"]]
        )

        # 10. Warns (Field: 'warning', 'description')
        self._render_section(
            title=self.labels["warns_header"],
            data=doc_data.get("warns", []),
            fields=['warning', 'description'],
            headers=[self.labels["col_warning"], self.labels["col_desc"]]
        )

        # 11. Warnings Section (Text Bebas)
        warnings_sec = doc_data.get("warnings_section", "")
        if warnings_sec:
            self.document.add_heading(self.labels["warnings_sec_header"], level=3)
            p = self.document.add_paragraph(warnings_sec)
            

        # 12. See Also
        self._render_section(
            title=self.labels["see_also_header"],
            data=doc_data.get("see_also", []),
            fields=['name', 'description'],
            headers=[self.labels["col_ref"], self.labels["col_desc"]]
        )

        # 13. Notes
        notes = doc_data.get("notes", "")
        if notes:
            self.document.add_heading(self.labels["notes_header"], level=3)
            self.document.add_paragraph(notes)

        # 14. Examples
        examples = doc_data.get("examples", "")
        if examples:
            self.document.add_heading(self.labels["examples_header"], level=3)
            self.document.add_paragraph(examples, style='CodeBlock')

        self.document.add_paragraph("_" * 50).alignment = WD_ALIGN_PARAGRAPH.CENTER

    def save(self, file_path: str):
        try:
            self.document.save(file_path)
            print(f"[DOCX] Berhasil disimpan: {file_path}")
        except Exception as e:
            print(f"[DOCX ERROR] Gagal simpan: {e}")
            
            
# --- FUNGSI HELPER UNTUK KONVERSI KE PDF (Windows Only) ---
def convert_docx_to_pdf(docx_path: str, pdf_path: str):
    """
    Mengonversi .docx ke .pdf menggunakan Microsoft Word (Windows).
    Membutuhkan: pip install docx2pdf
    """
    try:
        from docx2pdf import convert
        print("[PDF] Memulai konversi ke PDF...")
        convert(docx_path, pdf_path)
        print(f"[PDF] Berhasil dikonversi ke: {pdf_path}")
    except ImportError:
        print("[PDF ERROR] Library 'docx2pdf' belum diinstall. Jalankan: pip install docx2pdf")
    except Exception as e:
        print(f"[PDF ERROR] Konversi gagal: {e}")